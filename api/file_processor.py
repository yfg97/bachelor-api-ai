"""
File Processor für BA-Projekt: Steuerfahndung KI-Framework
Extrahiert Text aus verschiedenen Dateiformaten.

Unterstützte Formate:
- PDF (.pdf)
- Word (.docx, .doc)
- Text (.txt)
- CSV (.csv)
- E-Mail (.eml)
- Bulk Extractor Output (.txt mit speziellem Format)

Installation:
    pip install PyMuPDF python-docx --break-system-packages
"""

import os
import re
from pathlib import Path
from typing import Dict, Optional

# PDF-Extraktion
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  PyMuPDF nicht installiert. PDF-Support deaktiviert.")
    print("   Installiere mit: pip install PyMuPDF")

# Word-Dokumente
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx nicht installiert. DOCX-Support deaktiviert.")
    print("   Installiere mit: pip install python-docx")

# E-Mail-Parsing (eingebaut in Python)
import email
from email import policy
from email.parser import BytesParser


# =============================================================================
# HAUPTFUNKTIONEN
# =============================================================================

def process_file(file_path: str) -> Dict:
    """
    Hauptfunktion: Erkennt Dateityp und extrahiert Text.
    
    Args:
        file_path: Pfad zur Datei
        
    Returns:
        Dict mit:
        - success: bool
        - filename: str
        - filetype: str
        - text: str (extrahierter Text)
        - char_count: int
        - word_count: int
        - metadata: dict (optional, z.B. bei E-Mails)
        - error: str (nur bei Fehler)
    """
    path = Path(file_path)
    
    # Prüfe ob Datei existiert
    if not path.exists():
        return {
            "success": False,
            "error": f"Datei nicht gefunden: {file_path}"
        }
    
    suffix = path.suffix.lower()
    
    # Mapping: Dateiendung → Extraktionsfunktion
    extractors = {
        '.pdf': extract_from_pdf,
        '.docx': extract_from_docx,
        '.doc': extract_from_docx,
        '.txt': extract_from_txt,
        '.csv': extract_from_csv,
        '.eml': extract_from_eml,
        '.msg': extract_from_eml,  # Outlook
    }
    
    # Prüfe ob Format unterstützt
    if suffix not in extractors:
        return {
            "success": False,
            "error": f"Dateityp '{suffix}' nicht unterstützt",
            "supported_formats": list(extractors.keys())
        }
    
    # Extrahiere Text
    try:
        result = extractors[suffix](file_path)
        
        if result.get("success", False):
            text = result.get("text", "")
            result.update({
                "filename": path.name,
                "filetype": suffix,
                "char_count": len(text),
                "word_count": len(text.split()),
                "line_count": len(text.splitlines())
            })
        
        return result
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Fehler beim Verarbeiten: {str(e)}",
            "filename": path.name,
            "filetype": suffix
        }


# =============================================================================
# EINZELNE EXTRAKTOREN
# =============================================================================

def extract_from_pdf(file_path: str) -> Dict:
    """Extrahiert Text aus PDF-Datei."""
    if not PDF_AVAILABLE:
        return {
            "success": False,
            "error": "PyMuPDF nicht installiert. Führe aus: pip install PyMuPDF"
        }
    
    text_parts = []
    metadata = {}
    
    try:
        doc = fitz.open(file_path)
        
        # Metadaten extrahieren
        metadata = {
            "page_count": len(doc),
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "creation_date": doc.metadata.get("creationDate", "")
        }
        
        # Text aus allen Seiten
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Seite {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        return {
            "success": True,
            "text": "\n\n".join(text_parts),
            "metadata": metadata
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"PDF-Fehler: {str(e)}"
        }


def extract_from_docx(file_path: str) -> Dict:
    """Extrahiert Text aus Word-Dokument."""
    if not DOCX_AVAILABLE:
        return {
            "success": False,
            "error": "python-docx nicht installiert. Führe aus: pip install python-docx"
        }
    
    try:
        doc = Document(file_path)
        
        # Absätze extrahieren
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Tabellen extrahieren
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            tables_text.append("\n".join(table_rows))
        
        # Kombinieren
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n--- Tabellen ---\n" + "\n\n".join(tables_text)
        
        return {
            "success": True,
            "text": full_text,
            "metadata": {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"DOCX-Fehler: {str(e)}"
        }


def extract_from_txt(file_path: str) -> Dict:
    """Extrahiert Text aus TXT-Datei (mit Encoding-Erkennung)."""
    
    # Versuche verschiedene Encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            
            # Prüfe ob es Bulk Extractor Output ist
            if is_bulk_extractor_output(text):
                return parse_bulk_extractor(text, file_path)
            
            return {
                "success": True,
                "text": text,
                "metadata": {
                    "encoding": encoding
                }
            }
            
        except UnicodeDecodeError:
            continue
    
    return {
        "success": False,
        "error": "Konnte Datei-Encoding nicht erkennen"
    }


def extract_from_csv(file_path: str) -> Dict:
    """Extrahiert Text aus CSV-Datei."""
    import csv
    
    try:
        rows = []
        
        # Encoding erkennen
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    # Sniff delimiter
                    sample = f.read(4096)
                    f.seek(0)
                    
                    try:
                        dialect = csv.Sniffer().sniff(sample)
                    except csv.Error:
                        dialect = csv.excel
                    
                    reader = csv.reader(f, dialect)
                    for row in reader:
                        rows.append(" | ".join(row))
                
                break
            except UnicodeDecodeError:
                continue
        
        text = "\n".join(rows)
        
        return {
            "success": True,
            "text": text,
            "metadata": {
                "row_count": len(rows)
            }
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"CSV-Fehler: {str(e)}"
        }


def extract_from_eml(file_path: str) -> Dict:
    """Extrahiert Text und Metadaten aus E-Mail-Datei."""
    try:
        with open(file_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)
        
        # Metadaten
        metadata = {
            "from": str(msg.get('from', '')),
            "to": str(msg.get('to', '')),
            "cc": str(msg.get('cc', '')),
            "subject": str(msg.get('subject', '')),
            "date": str(msg.get('date', '')),
            "attachments": []
        }
        
        # Body extrahieren
        body_parts = []
        
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                
                if content_type == "text/plain":
                    try:
                        body_parts.append(part.get_content())
                    except Exception:
                        pass
                        
                # Anhänge notieren
                filename = part.get_filename()
                if filename:
                    metadata["attachments"].append(filename)
        else:
            try:
                body_parts.append(msg.get_content())
            except Exception:
                body_parts.append(str(msg.get_payload(decode=True)))
        
        # Formatierter Output
        header = f"""Von: {metadata['from']}
An: {metadata['to']}
CC: {metadata['cc']}
Betreff: {metadata['subject']}
Datum: {metadata['date']}
Anhänge: {', '.join(metadata['attachments']) if metadata['attachments'] else 'Keine'}
{'='*50}

"""
        
        body = "\n\n".join(body_parts)
        
        return {
            "success": True,
            "text": header + body,
            "metadata": metadata
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"E-Mail-Fehler: {str(e)}"
        }


# =============================================================================
# BULK EXTRACTOR SUPPORT (für forensische Daten)
# =============================================================================

def is_bulk_extractor_output(text: str) -> bool:
    """Prüft ob Text ein Bulk Extractor Output ist."""
    indicators = [
        "# BULK_EXTRACTOR",
        "# Feature-Recorder:",
        "# Filename:",
        ".E01"
    ]
    return any(indicator in text for indicator in indicators)


def parse_bulk_extractor(text: str, file_path: str) -> Dict:
    """
    Parst Bulk Extractor Output und strukturiert die Daten.
    Format: offset\tfeature_type\tcontext
    """
    lines = text.splitlines()
    
    # Header parsen
    metadata = {
        "bulk_extractor_version": "",
        "feature_recorder": "",
        "source_file": "",
        "feature_count": 0
    }
    
    features = []
    feature_types = {}
    
    for line in lines:
        # Header-Zeilen
        if line.startswith("# BULK_EXTRACTOR"):
            metadata["bulk_extractor_version"] = line.split(":")[-1].strip()
        elif line.startswith("# Feature-Recorder:"):
            metadata["feature_recorder"] = line.split(":")[-1].strip()
        elif line.startswith("# Filename:"):
            metadata["source_file"] = line.split(":")[-1].strip()
        elif line.startswith("#"):
            continue
        else:
            # Feature-Zeilen: offset\tfeature\tcontext
            parts = line.split("\t")
            if len(parts) >= 2:
                offset = parts[0]
                feature = parts[1] if len(parts) > 1 else ""
                context = parts[2] if len(parts) > 2 else ""
                
                features.append({
                    "offset": offset,
                    "feature": feature,
                    "context": context[:200]  # Kontext kürzen
                })
                
                # Feature-Typ zählen
                ftype = feature.split(":")[0] if ":" in feature else "unknown"
                feature_types[ftype] = feature_types.get(ftype, 0) + 1
    
    metadata["feature_count"] = len(features)
    metadata["feature_types"] = feature_types
    
    # Strukturierter Text-Output
    text_output = f"""BULK EXTRACTOR ANALYSE
{'='*50}
Quell-Image: {metadata['source_file']}
Feature-Recorder: {metadata['feature_recorder']}
Gefundene Features: {metadata['feature_count']}

Feature-Typen:
"""
    
    for ftype, count in sorted(feature_types.items(), key=lambda x: -x[1]):
        text_output += f"  - {ftype}: {count}\n"
    
    text_output += f"\n{'='*50}\nEXTRAHIERTE FEATURES:\n{'='*50}\n\n"
    
    # Erste 50 Features als Beispiel
    for i, feature in enumerate(features[:50]):
        text_output += f"[{feature['offset']}] {feature['feature']}\n"
        if feature['context']:
            text_output += f"    Kontext: {feature['context'][:100]}...\n"
        text_output += "\n"
    
    if len(features) > 50:
        text_output += f"\n... und {len(features) - 50} weitere Features\n"
    
    return {
        "success": True,
        "text": text_output,
        "metadata": metadata,
        "is_forensic_data": True,
        "raw_features": features[:100]  # Erste 100 für API
    }


# =============================================================================
# HILFSFUNKTIONEN
# =============================================================================

def get_supported_formats() -> list:
    """Gibt Liste unterstützter Formate zurück."""
    formats = ['.txt', '.csv', '.eml']
    
    if PDF_AVAILABLE:
        formats.append('.pdf')
    if DOCX_AVAILABLE:
        formats.extend(['.docx', '.doc'])
    
    return formats


def truncate_text(text: str, max_chars: int = 8000) -> str:
    """Kürzt Text für LLM-Verarbeitung."""
    if len(text) <= max_chars:
        return text
    
    # Behalte Anfang und Ende
    half = max_chars // 2
    return text[:half] + "\n\n[... Text gekürzt ...]\n\n" + text[-half:]


# =============================================================================
# TEST
# =============================================================================

if __name__ == "__main__":
    print("📁 File Processor Test")
    print("=" * 50)
    print(f"Unterstützte Formate: {get_supported_formats()}")
    print()
    
    # Test mit einer Datei wenn vorhanden
    test_files = [
        "test.txt",
        "test.pdf", 
        "test.docx"
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print(f"\nTeste: {test_file}")
            result = process_file(test_file)
            print(f"  Erfolg: {result.get('success')}")
            if result.get('success'):
                print(f"  Zeichen: {result.get('char_count')}")
                print(f"  Wörter: {result.get('word_count')}")
            else:
                print(f"  Fehler: {result.get('error')}")
